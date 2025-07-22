import os
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak,
    ListFlowable, ListItem, Table as RLTable, TableStyle, Flowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib import colors
from supabase import create_client, Client
from docx import Document
from docx.table import _Cell, Table
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph as DocxParagraph

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
SUPABASE_BUCKET = "pdfs"

def register_fonts():
    fonts_dir = os.path.join(os.path.dirname(__file__), "fonts")
    if not os.path.isdir(fonts_dir):
        print("⚠️ No fonts directory found.")
        return
    for filename in os.listdir(fonts_dir):
        if filename.lower().endswith(".ttf"):
            font_name = filename.rsplit(".", 1)[0]
            font_path = os.path.join(fonts_dir, filename)
            try:
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                print(f"✅ Registered font: {font_name}")
            except Exception as e:
                print(f"❌ Could not register font {font_name}: {e}")

def extract_book_title(docx_path):
    doc = Document(docx_path)
    # Try Heading 1 first
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading") and "1" in para.style.name and para.text.strip():
            return para.text.strip()
    # Fallback: first non-empty paragraph
    for para in doc.paragraphs:
        if para.text.strip():
            return para.text.strip()
    return "Untitled Manuscript"

def is_bullet_paragraph(para):
    style = para.style.name if hasattr(para.style, 'name') else ""
    if "List" in style or "Bullet" in style or "Number" in style:
        return True
    try:
        if para._element.numPr is not None:
            return True
    except Exception:
        pass
    return False

# Helper to handle drop caps in a very basic way
class DropCapParagraph(Flowable):
    def __init__(self, first_letter, rest_text, style, dropcap_size=30):
        Flowable.__init__(self)
        self.first_letter = first_letter
        self.rest_text = rest_text
        self.style = style
        self.dropcap_size = dropcap_size

    def wrap(self, availWidth, availHeight):
        return availWidth, self.dropcap_size + 4

    def draw(self):
        self.canv.setFont(self.style.fontName, self.dropcap_size)
        self.canv.drawString(0, 0, self.first_letter)
        self.canv.setFont(self.style.fontName, self.style.fontSize)
        self.canv.drawString(self.dropcap_size * 0.55, 0, self.rest_text)

def iter_block_items(parent):
    """
    Yield paragraphs and tables in document order.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        return
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield DocxParagraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def parse_docx_to_story(
    docx_path,
    styles,
    body_font="Roboto-Regular",
    heading_font="Roboto-Regular"
):
    doc = Document(docx_path)
    story = []
    list_buffer = []
    last_list_style = None
    used_title = False

    def flush_list():
        nonlocal list_buffer, last_list_style
        if list_buffer:
            story.append(ListFlowable(
                list_buffer,
                bulletType='bullet',
                leftIndent=18
            ))
            list_buffer = []
            last_list_style = None

    book_title = extract_book_title(docx_path)

    # Iterate paragraphs AND tables in order!
    for block in iter_block_items(doc):
        # Handle paragraph
        if isinstance(block, DocxParagraph):
            para = block
            text = ""
            for run in para.runs:
                t = run.text.replace('\n', ' ')
                if not t:
                    continue
                if run.bold:
                    t = f"<b>{t}</b>"
                if run.italic:
                    t = f"<i>{t}</i>"
                if run.underline:
                    t = f"<u>{t}</u>"
                text += t

            style = para.style.name if hasattr(para.style, 'name') else ""
            # Skip the real title if it was just used on the title page
            if not used_title and text.strip() == book_title:
                used_title = True
                continue
            # Drop Cap detection: if style includes "Drop Cap" or (customize for your needs)
            if "Drop Cap" in style:
                flush_list()
                if text.strip():
                    first_letter = text.strip()[0]
                    rest_text = text.strip()[1:]
                    # Use BookBody style for rest
                    story.append(DropCapParagraph(first_letter, rest_text, styles["BookBody"], dropcap_size=body_font_size*2))
                    story.append(Spacer(1, 6))
                continue
            if style.startswith('Heading'):
                flush_list()
                story.append(Spacer(1, 14))
                story.append(Paragraph(text, styles["BookHeading"]))
                story.append(Spacer(1, 10))
            elif is_bullet_paragraph(para):
                if last_list_style and style != last_list_style:
                    flush_list()
                list_buffer.append(ListItem(Paragraph(text, styles["BookBody"])))
                last_list_style = style
            elif text.strip() == "":
                flush_list()
                story.append(Spacer(1, 8))
            else:
                flush_list()
                story.append(Paragraph(text, styles["BookBody"]))
        # Handle tables in-order
        elif isinstance(block, Table):
            flush_list()
            data = []
            for row in block.rows:
                data.append([cell.text for cell in row.cells])
            t = RLTable(data, hAlign='LEFT')
            t.setStyle(TableStyle([
                ('FONTNAME', (0,0), (-1,-1), body_font),
                ('FONTSIZE', (0,0), (-1,-1), 12),
                ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ]))
            story.append(Spacer(1, 8))
            story.append(t)
            story.append(Spacer(1, 8))

    flush_list()
    return story

def generate_pdf(
    output_path: str,
    manuscript_file_path: str,   # Expects path to DOCX!
    heading_font: str = "Roboto-Regular",
    body_font: str = "Roboto-Regular",
    heading_size: float = 18.0,
    body_size: float = 12.0,
    trim_size: str = "6x9",
    bleed: bool = False,
    gutter: float = 0.25
):
    global body_font_size
    body_font_size = body_size

    trim_sizes = {
        "5x8": (5 * inch, 8 * inch),
        "5.5x8.5": (5.5 * inch, 8.5 * inch),
        "6x9": (6 * inch, 9 * inch),
        "7x10": (7 * inch, 10 * inch),
        "8.5x11": (8.5 * inch, 11 * inch),
    }
    page_size = trim_sizes.get(trim_size, (6 * inch, 9 * inch))
    width, height = page_size

    outer_margin = 0.75 * inch
    top_margin = 0.75 * inch
    bottom_margin = 0.75 * inch
    inner_margin = outer_margin + (gutter * inch)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="BookBody",
        fontName=body_font,
        fontSize=body_size,
        leading=body_size * 1.3,
        alignment=TA_JUSTIFY,
        firstLineIndent=body_size * 1.3,
        spaceAfter=body_size * 0.7,
    ))
    styles.add(ParagraphStyle(
        name="BookHeading",
        fontName=heading_font,
        fontSize=heading_size,
        alignment=TA_CENTER,
        spaceAfter=body_size * 1.2,
        leading=heading_size * 1.15,
    ))

    doc = SimpleDocTemplate(
        output_path,
        pagesize=(width, height),
        leftMargin=inner_margin,
        rightMargin=outer_margin,
        topMargin=top_margin,
        bottomMargin=bottom_margin,
    )

    # --- Use real book title from docx as title page ---
    book_title = extract_book_title(manuscript_file_path)
    story = []
    story.append(Spacer(1, height // 5))
    story.append(Paragraph(book_title, styles["BookHeading"]))
    story.append(PageBreak())

    # --- Parse DOCX file to story (with formatting, lists, tables, drop caps, order) ---
    story += parse_docx_to_story(manuscript_file_path, styles, body_font=body_font, heading_font=heading_font)

    def add_page_number(canvas, doc):
        page_num_text = "%d" % (doc.page)
        canvas.saveState()
        canvas.setFont(body_font, 10)
        canvas.drawCentredString(width / 2.0, 0.5 * inch, page_num_text)
        canvas.restoreState()

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    print(f"✅ PDF generated at: {output_path}")

def upload_pdf_to_supabase(pdf_path, pdf_filename):
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
    with open(pdf_path, "rb") as f:
        res = supabase.storage.from_(SUPABASE_BUCKET).upload(
            pdf_filename, f, {"content-type": "application/pdf", "upsert": "true"})
    public_url = f"{SUPABASE_URL}/storage/v1/object/public/{SUPABASE_BUCKET}/{pdf_filename}"
    return public_url
