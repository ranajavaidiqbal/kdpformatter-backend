import os
from docx import Document
from reportlab.platypus import Paragraph, Spacer, Table as RLTable, TableStyle, PageBreak
from reportlab.lib.styles import ParagraphStyle
from .bullets import parse_bullet_lists

def parse_docx_to_story(docx_path, styles):
    """
    Parses a DOCX file and returns a tuple:
      (story, headings)
    - story: List of ReportLab Flowables (Paragraphs, Tables, etc.)
    - headings: List of (heading_text, heading_level)
    """
    doc = Document(docx_path)
    story = []
    headings = []
    title_found = False

    # Gather all paragraphs (for list/bullet detection)
    all_paragraphs = list(doc.paragraphs)

    # Parse and add all bullet/numbered lists as flowables
    bullet_flowables = parse_bullet_lists(all_paragraphs, styles)
    story.extend(bullet_flowables)

    # Now process paragraphs for title, headings, normal paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Title detection (first non-empty para, before any heading)
        if not title_found and para.style.name.lower().startswith('title'):
            story.append(Paragraph(text, styles['BookHeading']))
            story.append(Spacer(1, 18))
            title_found = True
            continue

        # Heading detection
        if para.style.name.startswith('Heading'):
            # Extract heading level from style name, e.g., Heading1, Heading2
            try:
                level = int(para.style.name.replace('Heading', '').strip() or "1")
            except Exception:
                level = 1
            headings.append((text, level))
            story.append(Paragraph(text, styles['BookHeading']))
            story.append(Spacer(1, 12))
        # Already handled by bullets.py: skip any paragraph that's a list
        elif "list" in para.style.name.lower() or "bullet" in para.style.name.lower() or "number" in para.style.name.lower():
            continue
        # Body text (default)
        else:
            # Inline formatting (bold/italic) support
            run_fragments = []
            for run in para.runs:
                run_text = run.text.replace('\n', '')
                if not run_text:
                    continue
                if run.bold and run.italic:
                    run_fragments.append(f'<b><i>{run_text}</i></b>')
                elif run.bold:
                    run_fragments.append(f'<b>{run_text}</b>')
                elif run.italic:
                    run_fragments.append(f'<i>{run_text}</i>')
                else:
                    run_fragments.append(run_text)
            paragraph_text = ''.join(run_fragments)
            story.append(Paragraph(paragraph_text, styles['BookBody']))
        story.append(Spacer(1, 6))

    # Table handling (place all tables after paragraphs for now)
    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])
        rl_table = RLTable(data)
        rl_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), '#CCCCCC'),
            ('TEXTCOLOR', (0, 0), (-1, 0), '#000000'),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, '#000000'),
        ]))
        story.append(Spacer(1, 12))
        story.append(rl_table)
        story.append(Spacer(1, 12))

    return story, headings

def extract_book_title(docx_path):
    """Returns the first non-empty paragraph, or 'Untitled Book'."""
    doc = Document(docx_path)
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            return text
    return "Untitled Book"
